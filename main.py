import docx
import pygments
import os
import time

from docx.shared import Pt, Cm
from pygments.lexers import CppLexer
from pygments.formatters import RtfFormatter
from spire.doc import *
from spire.doc.common import *

USE_ADD_FILE = 1

def readFolder():
    print("Введите путь к папке с кодом для генерации Doc")
    folderName = input()
    return folderName

def readFile(file):
    f = open(file, "r", encoding='utf-8')
    str = f.read()
    lexer = CppLexer()
    formatter = RtfFormatter()
    highlighted_code = pygments.highlight(str, lexer, formatter)
    f2 = open("tmp.rtf", "w")
    f2.write(highlighted_code)
    f2.close()
    f.close()
    return highlighted_code

def insertOne(doc, file, i):
    code = readFile(file)
    file = file[file.rfind("\\")+1:]
    p = doc.add_paragraph(f"Листинг №{i}: \"Файл проекта {file}\"")
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)

    if not USE_ADD_FILE:
        p = doc.add_paragraph(code)
        run = p.runs[0]
        p.paragraph_format.left_indent = Cm(1.5)
        font = run.font
        font.name = 'Consolas'
        font.size = Pt(12)
    else:
        tmpDoc = Document()
        tmpDoc.LoadRtf("tmp.rtf")
        tmpDoc.SaveToFile("tmp2.docx")
        doc2 = docx.Document("tmp2.docx")
        doc2.save("test1")
        p_copyed = doc2.paragraphs
        p_local = doc.add_paragraph("")
        p_local.paragraph_format.left_indent = Cm(1.5)
        for p in p_copyed[1:]:
            for one_run in p.runs:
                if (one_run == p.runs[-1]):
                    local_run = p_local.add_run(one_run.text, one_run.style)
                else:
                    local_run = p_local.add_run(one_run.text, one_run.style)
                local_run.italic = one_run.italic
                local_run.bold = one_run.bold
                local_run.font.color.rgb = one_run.font.color.rgb
                font = local_run.font
                font.name = 'Consolas'
                font.size = Pt(12)
            p_local.add_run("\n")



def insertCode(doc, name):
    files = [each for each in os.listdir(name) if each.endswith('.cpp') or each.endswith('.h')]
    i = 1
    for file in files:
        if (name[-1] != '\\'):
            insertOne(doc, name + '\\' + file, i)
        else:
            insertOne(doc, name + file, i)
        i += 1
    try:
        os.remove('test1')
        os.remove('tmp.rtf')
        os.remove('tmp2.docx')
    except OSError:
        pass


if __name__ == '__main__':
    start_time = time.time()
    doc = docx.Document()
    name = readFolder()
    insertCode(doc, name)
    end_time = time.time()
    print(f"Перенос листингов в Docx завершен, он занял {end_time - start_time} секунд")
    doc.save('listings.docx')
